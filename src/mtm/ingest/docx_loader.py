"""Word document (.docx) ingestion."""

from datetime import datetime
from pathlib import Path
from typing import Optional
from uuid import uuid4

from docx import Document
from docx.document import Document as DocumentType

from mtm.models import Note


def parse_docx(file_path: str | Path) -> Note:
    """Parse Word document (.docx) file.

    Args:
        file_path: Path to .docx file

    Returns:
        Note object with parsed content and metadata

    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file cannot be parsed
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    # Load document
    doc: DocumentType = Document(str(file_path))

    # Extract text content
    content_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content_parts.append(paragraph.text.strip())

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                content_parts.append(row_text)

    content = "\n\n".join(content_parts)

    # Extract metadata from document properties
    core_props = doc.core_properties

    # Parse date
    date = None
    if core_props.created:
        date = core_props.created
    elif core_props.modified:
        date = core_props.modified
    else:
        date = datetime.fromtimestamp(file_path.stat().st_mtime)

    # Extract title
    meeting_title = (
        core_props.title
        or core_props.subject
        or _extract_title_from_content(content)
        or file_path.stem
    )

    # Extract metadata from custom properties or content
    # Try to find metadata in first paragraph or document properties
    metadata = {}
    if core_props.keywords:
        metadata["tags"] = core_props.keywords
    if core_props.comments:
        metadata["comments"] = core_props.comments
    if core_props.category:
        metadata["category"] = core_props.category

    # Try to extract structured metadata from content (first few paragraphs)
    content_lines = content.split("\n")[:10]
    extracted_metadata = _extract_metadata_from_content("\n".join(content_lines))

    # Normalize metadata
    project = extracted_metadata.get("project") or metadata.get("project") or "default"
    attendees = _normalize_list(extracted_metadata.get("attendees") or metadata.get("attendees") or [])
    roles = _normalize_list(extracted_metadata.get("roles") or metadata.get("roles") or [])
    tags = _normalize_list(extracted_metadata.get("tags") or metadata.get("tags") or [])

    # Create Note object
    note = Note(
        id=uuid4(),
        project=project,
        roles=roles,
        date=date,
        source_file=str(file_path),
        content=content,
        title=meeting_title,
        metadata={
            "attendees": ",".join(attendees) if attendees else "",
            "tags": ",".join(tags) if tags else "",
            "meeting": meeting_title,
            "author": core_props.author or "",
            "category": core_props.category or "",
            "comments": core_props.comments or "",
            **metadata,
        },
    )

    return note


def _extract_title_from_content(content: str) -> Optional[str]:
    """Extract title from content (first non-empty line or heading).

    Args:
        content: Document content

    Returns:
        Title string or None
    """
    lines = content.split("\n")
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if line and len(line) < 200:  # Reasonable title length
            # Remove common prefixes
            line = line.lstrip("#").strip()
            if line:
                return line
    return None


def _extract_metadata_from_content(content: str) -> dict[str, str | list[str]]:
    """Extract metadata patterns from content.

    Looks for patterns like:
    - Project: name
    - Attendees: person1, person2
    - Roles: role1, role2
    - Tags: tag1, tag2

    Args:
        content: Document content

    Returns:
        Dictionary of extracted metadata
    """
    import re

    metadata: dict[str, str | list[str]] = {}

    # Pattern matching for key-value pairs
    patterns = {
        "project": r"(?:project|project name|project_name)[:\s]+([^\n]+)",
        "attendees": r"(?:attendees|participants|people|present)[:\s]+([^\n]+)",
        "roles": r"(?:roles|role)[:\s]+([^\n]+)",
        "tags": r"(?:tags|tag)[:\s]+([^\n]+)",
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            if key in ["attendees", "roles", "tags"]:
                metadata[key] = [item.strip() for item in value.split(",") if item.strip()]
            else:
                metadata[key] = value

    return metadata


def _normalize_list(value: str | list | None) -> list[str]:
    """Normalize value to list of strings.

    Args:
        value: String, list, or None

    Returns:
        List of strings
    """
    if value is None:
        return []

    if isinstance(value, str):
        # Split by comma, semicolon, or newline
        return [item.strip() for item in value.split(",") if item.strip()]

    if isinstance(value, list):
        return [str(item).strip() for item in value if item]

    return []

