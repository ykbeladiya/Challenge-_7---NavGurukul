"""Generate realistic sample meeting notes for testing."""

import random
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from mtm.config import get_config


# Sample content templates
ONBOARDING_TOPICS = [
    "Welcome and introductions",
    "Company culture and values",
    "Development environment setup",
    "Code review process",
    "Git workflow and branching strategy",
    "Testing standards",
    "Documentation practices",
    "Team communication tools",
    "Performance expectations",
    "Career development paths",
]

DEPLOYMENT_TOPICS = [
    "Production deployment checklist",
    "Database migration strategy",
    "Rollback procedures",
    "Monitoring and alerting setup",
    "CI/CD pipeline configuration",
    "Load testing results",
    "Security audit findings",
    "Infrastructure scaling",
    "Disaster recovery plan",
    "Release notes preparation",
]

SUPPORT_TOPICS = [
    "Customer ticket escalation",
    "Bug triage and prioritization",
    "Knowledge base updates",
    "Common issues and solutions",
    "Customer feedback review",
    "Support metrics and KPIs",
    "Training new support staff",
    "Escalation procedures",
    "Documentation gaps",
    "Tool improvements",
]

ROLES = ["Engineer", "Support", "Product Manager"]

PROJECTS = {
    "Onboarding": {
        "topics": ONBOARDING_TOPICS,
        "roles": ["Engineer", "Product Manager"],
        "attendees": ["Alice", "Bob", "Charlie", "Diana", "Eve"],
    },
    "Deployment": {
        "topics": DEPLOYMENT_TOPICS,
        "roles": ["Engineer", "Product Manager"],
        "attendees": ["Frank", "Grace", "Henry", "Iris", "Jack"],
    },
    "Support": {
        "topics": SUPPORT_TOPICS,
        "roles": ["Support", "Engineer", "Product Manager"],
        "attendees": ["Karen", "Liam", "Mia", "Noah", "Olivia"],
    },
}


def generate_markdown_note(project: str, topic: str, date: datetime, index: int) -> str:
    """Generate a Markdown note with front matter."""
    project_data = PROJECTS[project]
    attendees_list = project_data["attendees"]
    roles_list = project_data["roles"]
    tags_list = ["urgent", "follow-up", "documentation", "training", "review"]
    
    attendees = random.sample(attendees_list, k=min(random.randint(2, 4), len(attendees_list)))
    roles = random.sample(roles_list, k=min(random.randint(1, len(roles_list)), len(roles_list)))
    tags = random.sample(tags_list, k=min(random.randint(1, 3), len(tags_list)))

    front_matter = f"""---
date: {date.strftime('%Y-%m-%d')}
meeting: {topic}
project: {project}
attendees:
"""
    for attendee in attendees:
        front_matter += f"  - {attendee}\n"
    front_matter += f"roles:\n"
    for role in roles:
        front_matter += f"  - {role}\n"
    front_matter += f"tags:\n"
    for tag in tags:
        front_matter += f"  - {tag}\n"
    front_matter += "---\n\n"

    content = f"""# {topic}

## Attendees
{', '.join(attendees)}

## Discussion

"""
    # Generate realistic discussion points
    discussion_points = [
        f"Discussed the importance of {topic.lower()} in the {project.lower()} process.",
        f"Reviewed current practices and identified areas for improvement.",
        f"Agreed on action items to implement changes.",
        f"Decided to schedule follow-up meeting next week.",
        f"Action: Update documentation with new guidelines.",
    ]

    discussion_sample_size = min(random.randint(3, 5), len(discussion_points))
    for point in random.sample(discussion_points, k=discussion_sample_size):
        content += f"- {point}\n"

    content += f"""
## Decisions

- **Decision:** Standardize the {topic.lower()} process across all teams.
- **Rationale:** This will improve consistency and reduce confusion.
- **Decision Maker:** {random.choice(attendees)}

## Action Items

- Action: {random.choice(attendees)} - Complete documentation update by {date + timedelta(days=7):%Y-%m-%d}
- Action: {random.choice(attendees)} - Review and approve changes by {date + timedelta(days=5):%Y-%m-%d}

## Next Steps

- Schedule follow-up meeting
- Review progress in next sprint
- Update team wiki with new information
"""

    return front_matter + content


def generate_text_note(project: str, topic: str, date: datetime, index: int) -> str:
    """Generate a plain text note."""
    project_data = PROJECTS[project]
    attendees_list = project_data["attendees"]
    attendees = random.sample(attendees_list, k=min(random.randint(2, 4), len(attendees_list)))

    content = f"""Meeting Notes: {topic}
Date: {date.strftime('%Y-%m-%d')}
Project: {project}
Attendees: {', '.join(attendees)}

Discussion:
"""
    discussion_points = [
        f"Reviewed {topic.lower()} procedures.",
        f"Identified key challenges and opportunities.",
        f"Agreed on next steps and timeline.",
    ]

    discussion_sample_size = min(random.randint(2, 4), len(discussion_points))
    for point in random.sample(discussion_points, k=discussion_sample_size):
        content += f"  - {point}\n"

    content += f"""
Decisions:
  - Standardized approach for {topic.lower()}
  - Timeline: {date + timedelta(days=14):%Y-%m-%d}

Action Items:
  - {random.choice(attendees)}: Complete task by {date + timedelta(days=7):%Y-%m-%d}
  - {random.choice(attendees)}: Review and provide feedback
"""

    return content


def generate_docx_note(project: str, topic: str, date: datetime, index: int, output_path: Path) -> None:
    """Generate a DOCX note."""
    if not DOCX_AVAILABLE:
        # Fallback to text format if python-docx not available
        generate_text_note(project, topic, date, index)
        return

    project_data = PROJECTS[project]
    attendees_list = project_data["attendees"]
    attendees = random.sample(attendees_list, k=min(random.randint(2, 4), len(attendees_list)))
    roles_list = project_data["roles"]
    roles = random.sample(roles_list, k=min(random.randint(1, len(roles_list)), len(roles_list)))

    doc = Document()
    
    # Title
    title = doc.add_heading(topic, 0)
    title_run = title.runs[0]
    title_run.font.size = Pt(18)

    # Metadata
    doc.add_paragraph(f"Date: {date.strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Project: {project}")
    doc.add_paragraph(f"Attendees: {', '.join(attendees)}")
    doc.add_paragraph(f"Roles: {', '.join(roles)}")

    doc.add_paragraph()  # Blank line

    # Discussion
    doc.add_heading("Discussion", level=1)
    discussion_points = [
        f"Discussed {topic.lower()} requirements and best practices.",
        f"Reviewed current implementation and identified gaps.",
        f"Agreed on action plan for improvements.",
    ]

    discussion_sample_size = min(random.randint(2, 3), len(discussion_points))
    for point in random.sample(discussion_points, k=discussion_sample_size):
        doc.add_paragraph(point, style="List Bullet")

    # Decisions
    doc.add_heading("Decisions", level=1)
    doc.add_paragraph(f"Decision: Implement standardized {topic.lower()} process")
    doc.add_paragraph(f"Rationale: Improves team efficiency and consistency")
    doc.add_paragraph(f"Decision Maker: {random.choice(attendees)}")

    # Action Items
    doc.add_heading("Action Items", level=1)
    doc.add_paragraph(f"Action: {random.choice(attendees)} - Complete task by {date + timedelta(days=7):%Y-%m-%d}")
    doc.add_paragraph(f"Action: {random.choice(attendees)} - Review documentation")

    doc.save(output_path)


def generate_pdf_note(project: str, topic: str, date: datetime, index: int, output_path: Path) -> None:
    """Generate a PDF note (as text file for now, since PDF generation requires additional libraries)."""
    # For simplicity, we'll create a text file that represents PDF content
    # In a real implementation, you'd use a PDF library like reportlab or pdfminer
    project_data = PROJECTS[project]
    attendees_list = project_data["attendees"]
    attendees = random.sample(attendees_list, k=min(random.randint(2, 4), len(attendees_list)))
    roles_list = project_data["roles"]
    roles = random.sample(roles_list, k=min(random.randint(1, len(roles_list)), len(roles_list)))
    
    content = f"""MEETING NOTES
==============

Topic: {topic}
Date: {date.strftime('%Y-%m-%d')}
Project: {project}
Attendees: {', '.join(attendees)}
Roles: {', '.join(roles)}

DISCUSSION:
-----------
- Reviewed {topic.lower()} procedures and best practices
- Identified key improvement areas
- Agreed on implementation timeline
- Discussed resource requirements

DECISIONS:
----------
- Standardize {topic.lower()} process
- Timeline: {date + timedelta(days=14):%Y-%m-%d}
- Decision Maker: {random.choice(attendees)}

ACTION ITEMS:
-------------
- {random.choice(attendees)}: Complete documentation update by {date + timedelta(days=7):%Y-%m-%d}
- {random.choice(attendees)}: Review and approve changes by {date + timedelta(days=5):%Y-%m-%d}
- Schedule follow-up meeting for next week
"""
    output_path.write_text(content, encoding="utf-8")


def generate_sample_notes(
    output_dir: Optional[str | Path] = None,
    num_notes: int = 20,
) -> list[Path]:
    """Generate sample meeting notes.

    Args:
        output_dir: Output directory (defaults to data/seed/notes)
        num_notes: Number of notes to generate

    Returns:
        List of generated file paths
    """
    config = get_config()

    if output_dir is None:
        output_dir = Path("data/seed/notes")
    else:
        output_dir = Path(output_dir)

    output_dir.mkdir(parents=True, exist_ok=True)

    generated_files: list[Path] = []
    projects = list(PROJECTS.keys())
    
    # Distribute notes across projects
    notes_per_project = num_notes // len(projects)
    remainder = num_notes % len(projects)

    # Generate notes for each project
    for project_idx, project in enumerate(projects):
        project_dir = output_dir / project
        project_dir.mkdir(parents=True, exist_ok=True)

        # Calculate number of notes for this project
        project_notes = notes_per_project + (1 if project_idx < remainder else 0)
        project_data = PROJECTS[project]

        # Generate notes with different formats
        # Distribute formats: 50% markdown, 25% text, 15% docx, 10% pdf
        format_counts = {
            ".md": max(1, project_notes // 2),
            ".txt": max(1, project_notes // 4),
            ".docx": max(1, project_notes // 6),
            ".pdf": max(0, project_notes - (project_notes // 2) - (project_notes // 4) - (project_notes // 6)),
        }
        
        # Ensure total matches project_notes
        total_formats = sum(format_counts.values())
        if total_formats < project_notes:
            format_counts[".md"] += project_notes - total_formats

        note_index = 0
        base_date = datetime.now() - timedelta(days=30)

        for fmt, count in format_counts.items():
            for i in range(count):
                if note_index >= project_notes:
                    break

                # Select topic
                topic = random.choice(project_data["topics"])
                
                # Generate date (spread over last 30 days)
                date = base_date + timedelta(days=random.randint(0, 30))

                # Generate filename (sanitize to avoid path issues)
                topic_slug = topic.lower().replace(" ", "-").replace(":", "").replace(",", "").replace("/", "-").replace("\\", "-")
                filename = f"{date.strftime('%Y%m%d')}_{topic_slug}{fmt}"
                file_path = project_dir / filename

                # Generate note based on format
                if fmt == ".md":
                    content = generate_markdown_note(project, topic, date, note_index)
                    file_path.write_text(content, encoding="utf-8")
                elif fmt == ".txt":
                    content = generate_text_note(project, topic, date, note_index)
                    file_path.write_text(content, encoding="utf-8")
                elif fmt == ".docx":
                    generate_docx_note(project, topic, date, note_index, file_path)
                elif fmt == ".pdf":
                    generate_pdf_note(project, topic, date, note_index, file_path)

                generated_files.append(file_path)
                note_index += 1

    return generated_files

